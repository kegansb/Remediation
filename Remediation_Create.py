from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
import json
from PIL import Image, ImageDraw, ImageFont

def create_default_logo(image_path):
    """Generate a simple logo if none exists."""
    if not os.path.exists(image_path):
        img = Image.new('RGB', (200, 50), color='white')
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("arial.ttf", 20)
        except:
            font = ImageFont.load_default()
        draw.text((10, 10), "CyberSec", fill='black', font=font)
        img.save(image_path, 'PNG')
        print(f"Generated default logo at {image_path}")
    else:
        print(f"Using existing image at {image_path}")

def create_checkbox_image(image_path):
    """Generate a small unchecked checkbox image."""
    if not os.path.exists(image_path):
        size = 20  # 20x20 pixels
        img = Image.new('RGB', (size, size), color='white')
        draw = ImageDraw.Draw(img)
        draw.rectangle([2, 2, size-3, size-3], outline='black', width=1)
        img.save(image_path, 'PNG')
        print(f"Generated checkbox image at {image_path}")
    else:
        print(f"Using existing checkbox image at {image_path}")

def load_remediation_templates(template_file="remediation_templates.json"):
    """Load remediation templates from an external JSON file."""
    try:
        if os.path.exists(template_file):
            with open(template_file, 'r') as f:
                templates = json.load(f)
            return templates
        else:
            print(f"Warning: {template_file} not found. Remediation actions will require manual input.")
            return {}
    except json.JSONDecodeError:
        print(f"Error: {template_file} is malformed. Remediation actions will require manual input.")
        return {}

def create_remediation_doc(filename, incident_data):
    # Create a new Document
    doc = Document()

    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Add image to header
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    logo_path = "logo.png"
    create_default_logo(logo_path)
    if os.path.exists(logo_path):
        run = header_paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        print(f"Warning: Image file '{logo_path}' not found. Header will be empty.")

    # Title
    title = doc.add_heading('Cyber Intrusion Remediation Document', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_heading('Metadata', level=2)
    doc.add_paragraph(f"Document ID: {incident_data['doc_id']}")
    doc.add_paragraph(f"Date Created: {incident_data['date_created']}")
    doc.add_paragraph(f"Prepared By: {incident_data['prepared_by']}")
    doc.add_paragraph(f"Response Team Lead: {incident_data['response_team_lead']}")
    doc.add_paragraph(f"Incident Date: {incident_data['incident_date']}")
    doc.add_paragraph(f"Last Updated: {incident_data['last_updated']}")
    doc.add_paragraph()

    # Section 1: Incident Overview
    doc.add_heading('1. Incident Overview', level=2)
    doc.add_paragraph('Purpose: Summarize the intrusion for context.', style='Normal').bold = True
    doc.add_paragraph('Details:', style='Normal').bold = True
    for key, value in incident_data['overview'].items():
        doc.add_paragraph(f"{key}: {value}", style='List Bullet')

    # Section 2: Intrusion Specifics
    doc.add_heading('2. Intrusion Specifics', level=2)
    doc.add_paragraph('Purpose: Provide detailed information about the intrusion to inform remediation.', style='Normal').bold = True
    doc.add_paragraph('Details:', style='Normal').bold = True
    for key, value in incident_data['specifics'].items():
        doc.add_paragraph(f"{key}: {value}", style='List Bullet')

    # Section 3: Remediation Plan
    doc.add_heading('3. Remediation Plan', level=2)
    doc.add_paragraph('Purpose: Outline specific actions to contain, eradicate, and recover from the intrusion based on its specifics.', style='Normal').bold = True

    # Generate checkbox image
    checkbox_path = "checkbox.png"
    create_checkbox_image(checkbox_path)

    # Subsections
    subsections = [
        ('3.1 Containment Actions', incident_data['remediation']['containment']),
        ('3.2 Eradication Actions', incident_data['remediation']['eradication']),
        ('3.3 Recovery Actions', incident_data['remediation']['recovery']),
        ('3.4 Preventive Measures', incident_data['remediation']['preventive'])
    ]
    for subsection_title, items in subsections:
        doc.add_heading(subsection_title, level=3)
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run()
            if os.path.exists(checkbox_path):
                run.add_picture(checkbox_path, width=Inches(0.15))
                run.add_text(f" {item}")
            else:
                run.add_text(f"[ ] {item}")

    doc.add_paragraph('Notes:', style='Normal').bold = True
    doc.add_paragraph(incident_data['remediation']['notes'])

    # Section 4: Post-Incident Actions
    doc.add_heading('4. Post-Incident Actions', level=2)
    doc.add_paragraph('Purpose: Ensure lessons learned and compliance.', style='Normal').bold = True
    doc.add_paragraph('Details:', style='Normal').bold = True
    for key, value in incident_data['post_incident'].items():
        doc.add_paragraph(f"{key}: {value}", style='List Bullet')

    # Section 5: Stakeholders
    doc.add_heading('5. Stakeholders', level=2)
    doc.add_paragraph('Purpose: Identify key contacts for accountability.', style='Normal').bold = True
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Details'
    for role, details in incident_data['stakeholders']:
        row_cells = table.add_row().cells
        row_cells[0].text = role
        row_cells[1].text = details

    # Section 6: Appendices
    doc.add_heading('6. Appendices', level=2)
    for key, value in incident_data['appendices'].items():
        doc.add_paragraph(f"{key}: {value}", style='List Bullet')

    # Save the document
    doc.save(filename)
    print(f"Document saved as {filename}")

def get_user_input():
    print("Enter incident details (press Enter to use defaults where prompted):")
    
    # Get current date as default
    current_date = datetime.now().strftime("%Y-%m-%d")

    # Load remediation templates from external file
    remediation_templates = load_remediation_templates()
    
    # List valid incident types
    if remediation_templates:
        valid_types = list(remediation_templates.keys())
        print(f"Valid Incident Types: {', '.join(valid_types)}")

    # Metadata
    doc_id = input("Document ID (e.g., INC-2025-04-14-001): ") or "INC-YYYY-MM-DD-001"
    filename = f"{doc_id.replace(':', '_').replace('/', '_')}.docx"
    
    # Check for existing file
    while os.path.exists(filename):
        print(f"Warning: {filename} already exists.")
        overwrite = input("Overwrite? (y/n): ").lower()
        if overwrite == 'y':
            break
        new_id = input("Enter new Document ID (e.g., INC-2025-04-14-002): ") or "INC-YYYY-MM-DD-002"
        doc_id = new_id
        filename = f"{new_id.replace(':', '_').replace('/', '_')}.docx"

    date_created = input(f"Date Created (default: {current_date}): ") or current_date
    prepared_by = input("Prepared By (e.g., John Doe/Team Name): ") or "[Your Name/Team Name]"
    response_team_lead = input("Response Team Lead (e.g., Jane Smith): ") or "[Response Team Lead]"
    incident_date = input(f"Incident Date (default: {current_date}): ") or current_date
    last_updated = input(f"Last Updated (default: {current_date}): ") or current_date

    # Incident Overview
    incident_type = input("Incident Type (e.g., Malware, Phishing, Ransomware): ") or "[e.g., Malware, Phishing]"
    # Standardize incident_type for case-insensitive matching
    incident_type_standardized = incident_type.lower().capitalize()
    incident_severity = input("Incident Severity (e.g., Low, Medium, High, Critical): ") or "[e.g., Low, Medium, High, Critical]"
    date_time_detected = input("Date/Time Detected (e.g., 2025-04-13 14:30): ") or "[Insert Date/Time]"
    affected_systems = input("Affected Systems/Assets (e.g., Servers, Endpoints): ") or "[e.g., Servers, Endpoints]"
    impact_summary = input("Impact Summary (e.g., Data Breach, Service Disruption): ") or "[e.g., Data Breach, Service Disruption]"

    # Intrusion Specifics
    attack_vector = input("Attack Vector (e.g., Exploited Vulnerability, Stolen Credentials): ") or "[e.g., Exploited Vulnerability, Stolen Credentials]"
    iocs = input("Indicators of Compromise (e.g., Malicious IPs, Hashes): ") or "[e.g., Malicious IPs, Hashes]"
    scope = input("Scope of Compromise (e.g., 10 Devices, 2 Databases): ") or "[e.g., Number of Affected Devices]"
    root_cause = input("Root Cause (if known, e.g., Unpatched Software): ") or "[e.g., Unpatched Software]"

    # Remediation Actions
    if remediation_templates and incident_type_standardized in remediation_templates:
        print(f"Auto-populating remediation actions for {incident_type_standardized} incident.")
        containment = remediation_templates[incident_type_standardized]['containment']
        eradication = remediation_templates[incident_type_standardized]['eradication']
        recovery = remediation_templates[incident_type_standardized]['recovery']
        preventive = remediation_templates[incident_type_standardized]['preventive']
    else:
        print(f"Incident Type '{incident_type}' not in predefined list or templates unavailable. Enter remediation actions manually.")
        print("Enter Containment Actions (one per line, type 'done' when finished):")
        containment = []
        while True:
            action = input("> ")
            if action.lower() == 'done':
                break
            if action:
                containment.append(action)
        if not containment:
            containment = ["Short-Term Containment: [e.g., Isolate affected systems]", "Long-Term Containment: [e.g., Deploy network segmentation]"]

        print("Enter Eradication Actions (one per line, type 'done' when finished):")
        eradication = []
        while True:
            action = input("> ")
            if action.lower() == 'done':
                break
            if action:
                eradication.append(action)
        if not eradication:
            eradication = ["Remove Malicious Artifacts: [e.g., Delete malware]", "Patch Vulnerabilities: [e.g., Apply software updates]"]

        print("Enter Recovery Actions (one per line, type 'done' when finished):")
        recovery = []
        while True:
            action = input("> ")
            if action.lower() == 'done':
                break
            if action:
                recovery.append(action)
        if not recovery:
            recovery = ["Restore Systems/Services: [e.g., Rebuild servers]", "Validate Integrity: [e.g., Verify no residual threats]"]

        print("Enter Preventive Measures (one per line, type 'done' when finished):")
        preventive = []
        while True:
            action = input("> ")
            if action.lower() == 'done':
                break
            if action:
                preventive.append(action)
        if not preventive:
            preventive = ["Based on Intrusion Specifics: [e.g., If phishing, enhance email filters]", "General Hardening: [e.g., Update firewall rules]"]

    # Post-Incident Actions
    regulatory_notification = input("Regulatory Notification Required (e.g., Yes/No): ") or "[e.g., Yes/No]"

    # Construct incident_data dictionary
    incident_data = {
        'doc_id': doc_id,
        'date_created': date_created,
        'prepared_by': prepared_by,
        'response_team_lead': response_team_lead,
        'incident_date': incident_date,
        'last_updated': last_updated,
        'overview': {
            'Incident Type': incident_type,
            'Incident Severity': incident_severity,
            'Date/Time Detected': date_time_detected,
            'Affected Systems/Assets': affected_systems,
            'Impact Summary': impact_summary,
            'Initial Detection Method': '[e.g., IDS Alert, User Report]'
        },
        'specifics': {
            'Attack Vector': attack_vector,
            'Indicators of Compromise (IoCs)': iocs,
            'Scope of Compromise': scope,
            'Root Cause (if known)': root_cause,
            'Threat Actor (if identified)': '[e.g., Known Group, Unknown]',
            'Evidence Collected': '[e.g., Logs, Memory Dumps]'
        },
        'remediation': {
            'containment': containment,
            'eradication': eradication,
            'recovery': recovery,
            'preventive': preventive,
            'notes': ('Tailor this section to the intrusion’s root cause and attack vector. '
                      'For example, a ransomware incident may prioritize backup restoration, '
                      'while a credential theft incident may focus on MFA enforcement.')
        },
        'post_incident': {
            'Incident Review Date': '[Schedule Date]',
            'Lessons Learned': '[e.g., Identified gaps in monitoring]',
            'Policy Updates': '[e.g., Revise incident response plan]',
            'Regulatory Notification Required': regulatory_notification,
            'Reporting Requirements': '[e.g., Notify regulators]',
            'Documentation Status': '[e.g., Final report archived]'
        },
        'stakeholders': [
            ('Incident Lead', '[Name, Role, Contact]'),
            ('IT/Security Team', '[Name(s), Role(s)]'),
            ('External Partners', '[e.g., Forensics Firm, Law Enforcement]'),
            ('Approver', '[Name, Role]')
        ],
        'appendices': {
            'Logs/Reports': '[Reference attached evidence or logs]',
            'Timeline of Events': '[Detailed chronology of incident and response]',
            'Additional Notes': '[Any other relevant information]'
        }
    }
    return incident_data, filename

# Run the script
if __name__ == "__main__":
    incident_data, filename = get_user_input()
    create_remediation_doc(filename, incident_data)
